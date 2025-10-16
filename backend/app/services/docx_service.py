from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import logging
from app.core.cache import cached
from app.core.constants import CacheConstants

logger = logging.getLogger(__name__)

class DOCXService:
    """Service for generating DOCX resumes"""
    
    @staticmethod
    @cached(CacheConstants.TEMPLATE_CACHE_TTL)
    def generate_resume_docx(resume_data: dict) -> bytes:
        """Generate a DOCX file from resume data"""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Personal Info Header
        personal_info = resume_data.get("personal_info", {})
        name = doc.add_heading(personal_info.get("full_name", ""), level=1)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact_parts = []
        if personal_info.get("email"):
            contact_parts.append(personal_info["email"])
        if personal_info.get("phone"):
            contact_parts.append(personal_info["phone"])
        if personal_info.get("location"):
            contact_parts.append(personal_info["location"])
        
        if contact_parts:
            contact = doc.add_paragraph(" | ".join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if personal_info.get("linkedin"):
            linkedin = doc.add_paragraph(personal_info["linkedin"])
            linkedin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if personal_info.get("summary"):
            doc.add_heading("Professional Summary", level=2)
            doc.add_paragraph(personal_info["summary"])
        
        # Experience
        experience = resume_data.get("experience", [])
        if experience:
            doc.add_heading("Experience", level=2)
            for exp in experience:
                # Position and Company
                p = doc.add_paragraph()
                p.add_run(exp.get("position", "")).bold = True
                p.add_run(f" | {exp.get('company', '')}")
                
                # Dates
                dates = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                date_p = doc.add_paragraph(dates)
                date_p.runs[0].italic = True
                
                # Description
                if exp.get("description"):
                    doc.add_paragraph(exp["description"])
                
                # Achievements
                if exp.get("achievements"):
                    for achievement in exp["achievements"]:
                        doc.add_paragraph(achievement, style='List Bullet')
        
        # Education
        education = resume_data.get("education", [])
        if education:
            doc.add_heading("Education", level=2)
            for edu in education:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
                doc.add_paragraph(edu.get("institution", ""))
                dates = f"{edu.get('start_date', '')} - {edu.get('end_date', 'Present')}"
                date_p = doc.add_paragraph(dates)
                date_p.runs[0].italic = True
                if edu.get("gpa"):
                    doc.add_paragraph(f"GPA: {edu['gpa']}")
        
        # Skills
        skills = resume_data.get("skills", [])
        if skills:
            doc.add_heading("Skills", level=2)
            skill_names = [s.get("name", "") for s in skills]
            doc.add_paragraph(", ".join(skill_names))
        
        # Projects
        projects = resume_data.get("projects", [])
        if projects:
            doc.add_heading("Projects", level=2)
            for project in projects:
                p = doc.add_paragraph()
                p.add_run(project.get("name", "")).bold = True
                doc.add_paragraph(project.get("description", ""))
                if project.get("technologies"):
                    tech = ", ".join(project["technologies"])
                    doc.add_paragraph(f"Technologies: {tech}")
        
        # Certifications
        certifications = resume_data.get("certifications", [])
        if certifications:
            doc.add_heading("Certifications", level=2)
            for cert in certifications:
                p = doc.add_paragraph()
                p.add_run(cert.get("name", "")).bold = True
                p.add_run(f" | {cert.get('issuer', '')}")
                doc.add_paragraph(cert.get("date", ""))
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
